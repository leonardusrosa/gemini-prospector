#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera contrato .docx protegido para envio PRIVADO ao cliente.

Uso:
    python3 gerar-docx.py dados.json saida.docx

IMPORTANTE:
- O contrato contém dados pessoais e condições comerciais. Não publicar em rota pública.
- O domínio é de responsabilidade do cliente.
- Hospedagem e manutenção são conceitos separados.
- O editor cobre alterações simples suportadas pela ferramenta; mudanças estruturais podem ser orçadas separadamente.

Campos que o CLIENTE pode preencher nas regiões editáveis:
CPF_CNPJ_CLIENTE e ENDERECO_CLIENTE quando vierem como "(preencher)", data e assinatura.
"""
import json
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

if len(sys.argv) != 3:
    raise SystemExit("Uso: python3 gerar-docx.py dados.json saida.docx")

d = json.load(open(sys.argv[1], encoding="utf-8"))
PID = [100]

REQUIRED = [
    "NOME_CLIENTE", "CPF_CNPJ_CLIENTE_LABEL", "CIDADE_UF_CLIENTE",
    "NOME_PRESTADOR", "CPF_CNPJ_PRESTADOR_LABEL", "CPF_CNPJ_PRESTADOR",
    "ENDERECO_PRESTADOR", "CIDADE_UF_PRESTADOR", "TEXTO_OBJETO",
    "URL_PUBLICADA", "VALOR", "VALOR_EXTENSO", "FORMA_PAGAMENTO",
    "PRAZO_ENTREGA", "RODADAS_AJUSTES", "TEXTO_HOSPEDAGEM",
    "CIDADE_FORO", "CIDADE_ASSINATURA"
]
missing = [k for k in REQUIRED if not str(d.get(k, "")).strip()]
if missing:
    raise ValueError("Campos obrigatórios ausentes no contrato: " + ", ".join(missing))

is_dry_run = bool(d.get("dry_run") or d.get("DRY_RUN") or d.get("modo") == "dry_run" or d.get("is_dry_run"))

if not is_dry_run:
    provider_fields = [
        ("NOME_PRESTADOR", d.get("NOME_PRESTADOR")),
        ("CPF_CNPJ_PRESTADOR", d.get("CPF_CNPJ_PRESTADOR")),
        ("ENDERECO_PRESTADOR", d.get("ENDERECO_PRESTADOR")),
        ("CIDADE_UF_PRESTADOR", d.get("CIDADE_UF_PRESTADOR")),
    ]
    for field_name, val in provider_fields:
        val_str = str(val or "").strip()
        if not val_str or any(marker in val_str.lower() for marker in ["preencher", "fictício", "ficticio", "000.000.000-00", "00.000.000/0000-00"]):
            raise ValueError(
                f"Contrato real bloqueado: campo obrigatório do prestador '{field_name}' ausente ou contém placeholder."
            )

# Evita reintroduzir a lógica comercial antiga em que hospedagem e manutenção eram sinônimos.
legacy_phrases = [
    "manutenção mensal da página (hospedagem, pequenas atualizações de texto/imagens e suporte)",
    "a contratação e renovação de hospedagem e domínio próprios são de responsabilidade do contratante"
]
hosting_lower = str(d.get("TEXTO_HOSPEDAGEM", "")).lower()
for phrase in legacy_phrases:
    if phrase in hosting_lower:
        raise ValueError("TEXTO_HOSPEDAGEM contém redação comercial antiga e incompatível com a oferta atual.")

if d.get("MANUTENCAO"):
    if not str(d.get("VALOR_MANUTENCAO", "")).strip():
        raise ValueError("VALOR_MANUTENCAO é obrigatório quando MANUTENCAO=true.")
    if not str(d.get("TEXTO_MANUTENCAO", "")).strip():
        raise ValueError("TEXTO_MANUTENCAO é obrigatório quando MANUTENCAO=true; não presumir escopo recorrente.")


def par(doc, texto="", bold=False, center=False, size=11, antes=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    if texto:
        r = p.add_run(texto)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = "Georgia"
    return p


def run(p, texto, bold=False, size=11):
    r = p.add_run(texto)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Georgia"
    return r


def editavel(p, texto):
    """Insere um trecho que o cliente pode editar (permStart/permEnd, grupo everyone)."""
    PID[0] += 1
    pid = str(PID[0])
    ps = OxmlElement("w:permStart")
    ps.set(qn("w:id"), pid)
    ps.set(qn("w:edGrp"), "everyone")
    p._p.append(ps)
    r = run(p, texto)
    r.font.highlight_color = 7
    pe = OxmlElement("w:permEnd")
    pe.set(qn("w:id"), pid)
    p._p.append(pe)


def campo(p, valor, rotulo):
    if "preencher" in (valor or "").lower() or not valor:
        editavel(p, " [" + rotulo + ": preencher aqui] ")
    else:
        run(p, valor)


def clausula(doc, n, titulo, texto):
    par(doc, "Cláusula %sª - %s" % (n, titulo), bold=True, antes=12)
    par(doc, texto)


doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.2)
    s.left_margin = s.right_margin = Cm(2.2)

par(doc, "CONTRATO DE PRESTAÇÃO DE SERVIÇOS", bold=True, center=True, size=13)
par(doc, "CRIAÇÃO E PUBLICAÇÃO DE PÁGINA NA INTERNET", bold=True, center=True, size=11)

p = par(doc)
run(p, "CONTRATANTE: ", bold=True)
run(p, d["NOME_CLIENTE"] + ", " + d["CPF_CNPJ_CLIENTE_LABEL"] + " nº ")
campo(p, d.get("CPF_CNPJ_CLIENTE"), "CPF/CNPJ")
run(p, ", com endereço em ")
campo(p, d.get("ENDERECO_CLIENTE"), "endereço")
run(p, ", " + d["CIDADE_UF_CLIENTE"] + ".")

p = par(doc)
run(p, "CONTRATADO(A): ", bold=True)
run(
    p,
    "%s, %s nº %s, com endereço em %s, %s."
    % (
        d["NOME_PRESTADOR"], d["CPF_CNPJ_PRESTADOR_LABEL"],
        d["CPF_CNPJ_PRESTADOR"], d["ENDERECO_PRESTADOR"], d["CIDADE_UF_PRESTADOR"]
    )
)

par(doc, "As partes acima identificadas celebram o presente contrato de prestação de serviços, que se regerá pelas cláusulas seguintes.")

clausula(
    doc, 1, "Do objeto",
    "%s, incluindo adaptação para dispositivos móveis, preparação para publicação e disponibilização de editor de conteúdo nos limites técnicos da solução entregue. A página será publicada no endereço %s, após as aprovações necessárias do CONTRATANTE."
    % (d["TEXTO_OBJETO"], d["URL_PUBLICADA"])
)
clausula(
    doc, 2, "Do valor e forma de pagamento",
    "Pelos serviços descritos na Cláusula 1ª, o CONTRATANTE pagará ao CONTRATADO(A) o valor total de R$ %s (%s), na seguinte forma: %s."
    % (d["VALOR"], d["VALOR_EXTENSO"], d["FORMA_PAGAMENTO"])
)
clausula(
    doc, 3, "Do prazo de entrega",
    "A página em sua versão final será entregue e publicada em até %s a contar da assinatura deste contrato e do fornecimento, pelo CONTRATANTE, dos materiais e aprovações necessários. Está incluída %s rodada(s) de ajustes de texto e imagens após a entrega."
    % (d["PRAZO_ENTREGA"], d["RODADAS_AJUSTES"])
)

n = 4
if d.get("MANUTENCAO"):
    clausula(
        doc, 4, "Do serviço adicional de manutenção",
        "O CONTRATANTE contrata serviço adicional de manutenção mensal pelo valor de R$ %s mensais, com o seguinte escopo expressamente acordado: %s. Itens não descritos neste escopo não são considerados incluídos automaticamente."
        % (d["VALOR_MANUTENCAO"], d["TEXTO_MANUTENCAO"])
    )
    n = 5

clausula(
    doc, n, "Do conteúdo e responsabilidades",
    "O CONTRATANTE declara ser titular ou possuir autorização de uso de todos os textos, imagens, logotipo e informações fornecidos, responsabilizando-se pela veracidade das informações profissionais divulgadas. O CONTRATADO(A) compromete-se a não inserir na página informações não fornecidas, não verificadas ou não aprovadas pelo CONTRATANTE."
)
clausula(doc, n + 1, "Do domínio e da hospedagem", d["TEXTO_HOSPEDAGEM"])
clausula(
    doc, n + 2, "Do editor de conteúdo e alterações futuras",
    "O CONTRATANTE poderá utilizar o editor de conteúdo disponibilizado para realizar alterações simples suportadas pela ferramenta, como textos, imagens, contatos, WhatsApp e links, sem cobrança por cada edição realizada diretamente pelo painel. Alterações estruturais ou complexas, incluindo novas páginas ou seções relevantes, novas integrações, novas funcionalidades, fluxos especiais ou redesign estrutural, não integram automaticamente o escopo original e poderão ser objeto de orçamento separado, sempre submetido à aprovação do CONTRATANTE antes da execução."
)
clausula(
    doc, n + 3, "Da rescisão",
    "Este contrato poderá ser rescindido por qualquer das partes mediante comunicação por escrito. Em caso de rescisão pelo CONTRATANTE após o início dos trabalhos, será devido o valor proporcional aos serviços já executados. Serviços adicionais recorrentes, quando expressamente contratados, observarão as condições específicas acordadas entre as partes."
)
clausula(
    doc, n + 4, "Do foro",
    "Fica eleito o foro da comarca de %s para dirimir quaisquer controvérsias oriundas deste contrato." % d["CIDADE_FORO"]
)

p = par(doc, antes=18)
run(p, d["CIDADE_ASSINATURA"] + ", ")
if str(d.get("DATA_EXTENSO", "")).strip():
    run(p, d["DATA_EXTENSO"])
else:
    editavel(p, " [data] ")
run(p, ".")

par(doc, "", antes=24)
p = par(doc, antes=18)
run(p, "__________________________________________")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = par(doc, antes=0)
run(p, d["NOME_CLIENTE"] + " - Contratante  ", bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
editavel(p, " [assine aqui] ")
p = par(doc, antes=18)
run(p, "__________________________________________")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = par(doc, antes=0, center=True)
run(p, d["NOME_PRESTADOR"] + " - Contratado(a)", bold=True)

par(
    doc,
    "Este documento é uma minuta base gerada automaticamente. Recomenda-se a revisão por profissional jurídico antes da assinatura.",
    size=8,
    antes=20
)

# Proteção: somente leitura, exceto as regiões permitidas acima.
dp = OxmlElement("w:documentProtection")
dp.set(qn("w:edit"), "readOnly")
dp.set(qn("w:enforcement"), "1")
doc.settings.element.append(dp)

doc.save(sys.argv[2])
print("docx gerado:", sys.argv[2])
